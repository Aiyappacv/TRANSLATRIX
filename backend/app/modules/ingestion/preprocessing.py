from __future__ import annotations

import hashlib
import json
import mimetypes
import os
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional
from uuid import UUID

import structlog
from sqlalchemy.orm import Session

from app.config import settings
from app.modules.ingestion.tiered_storage import (
    LakeTier,
    ProcessingStep,
    TieredStorageManager,
)

logger = structlog.get_logger(__name__)


class PreprocessingPipeline:
    def __init__(self, db: Session, storage_manager: TieredStorageManager):
        self.db = db
        self.storage = storage_manager

    def _now(self) -> str:
        return datetime.now(timezone.utc).isoformat()

    async def run(
        self,
        file_content: bytes,
        tenant_id: UUID,
        company_id: UUID,
        file_id: UUID,
        filename: str,
        content_type: str,
        checksum: str,
    ) -> dict[str, Any]:
        started = datetime.now(timezone.utc)
        metadata: dict[str, Any] = {}
        self.storage.log_audit(file_id, ProcessingStep.PREPROCESSING, "started", message=f"Preprocessing started for {filename}")

        metadata = self._generate_metadata(file_content, filename, content_type, checksum)
        metadata["preprocessingStartedAt"] = self._now()

        parsed = self._parse_document(file_content, filename, content_type)
        metadata["pageCount"] = parsed.get("pageCount", 0)
        metadata["wordCount"] = parsed.get("wordCount", 0)
        metadata["hasText"] = parsed.get("hasText", False)
        metadata["languageHint"] = parsed.get("languageHint")
        metadata["parsedStructure"] = parsed.get("structure", "unknown")

        raw_record = await self.storage.store_raw(
            file_content=file_content,
            tenant_id=tenant_id,
            company_id=company_id,
            file_id=file_id,
            filename=filename,
            content_type=content_type,
            checksum=checksum,
            metadata=metadata,
        )

        processed_data = {
            "documentId": str(file_id),
            "filename": filename,
            "contentType": content_type,
            "checksum": checksum,
            "metadata": metadata,
            "parsedContent": parsed.get("textContent", ""),
            "pages": parsed.get("pages", []),
            "tables": parsed.get("tables", []),
        }
        processed_record = await self.storage.store_processed(
            data=processed_data,
            tenant_id=tenant_id,
            company_id=company_id,
            file_id=file_id,
            filename=filename,
        )

        features = self._extract_features(metadata, parsed, filename)
        curated_record = await self.storage.store_curated(
            features=features,
            tenant_id=tenant_id,
            company_id=company_id,
            file_id=file_id,
            filename=filename,
        )

        duration_ms = int((datetime.now(timezone.utc) - started).total_seconds() * 1000)
        self.storage.log_audit(
            file_id,
            ProcessingStep.PREPROCESSED,
            "completed",
            message=f"Preprocessing completed for {filename}",
            duration_ms=duration_ms,
            metadata_json={
                "rawStorageKey": raw_record.storage_key,
                "processedStorageKey": processed_record.storage_key,
                "curatedStorageKey": curated_record.storage_key,
                "pageCount": metadata.get("pageCount"),
                "wordCount": metadata.get("wordCount"),
            },
        )

        return {
            "fileId": str(file_id),
            "filename": filename,
            "metadata": metadata,
            "tiers": {
                "raw": {"storageKey": raw_record.storage_key, "id": str(raw_record.id)},
                "processed": {"storageKey": processed_record.storage_key, "id": str(processed_record.id)},
                "curated": {"storageKey": curated_record.storage_key, "id": str(curated_record.id)},
            },
            "durationMs": duration_ms,
        }

    def _generate_metadata(
        self, content: bytes, filename: str, content_type: str, checksum: str
    ) -> dict[str, Any]:
        ext = Path(filename).suffix.lower()
        mime_type = content_type or mimetypes.guess_type(filename)[0] or "application/octet-stream"
        size_bytes = len(content)

        return {
            "filename": filename,
            "extension": ext,
            "mimeType": mime_type,
            "sizeBytes": size_bytes,
            "checksum": checksum,
            "checksumAlgorithm": "sha256",
            "isImage": mime_type.startswith("image/"),
            "isPdf": ext == ".pdf" or mime_type == "application/pdf",
            "isSpreadsheet": ext in {".csv", ".xlsx", ".xls", ".xlsm"},
            "isDocument": ext in {".docx", ".txt", ".json", ".xml", ".md"},
            "detectedAt": self._now(),
        }

    def _parse_document(self, content: bytes, filename: str, content_type: str) -> dict[str, Any]:
        ext = Path(filename).suffix.lower()
        result: dict[str, Any] = {
            "pageCount": 0,
            "wordCount": 0,
            "hasText": False,
            "languageHint": None,
            "structure": "unknown",
            "textContent": "",
            "pages": [],
            "tables": [],
        }

        raw_text = ""
        try:
            if ext == ".pdf":
                result["structure"] = "pdf"
                result["pageCount"] = self._count_pdf_pages(content)
                raw_text = self._extract_pdf_text(content)
            elif content_type.startswith("image/") or ext in {".png", ".jpg", ".jpeg", ".webp", ".tiff", ".tif", ".bmp"}:
                result["structure"] = "image"
                result["pageCount"] = 1
            elif ext in {".csv", ".xlsx", ".xls", ".xlsm"}:
                result["structure"] = "spreadsheet"
                result["pageCount"] = 1
                raw_text = self._extract_spreadsheet_text(content, ext)
            elif ext == ".docx":
                result["structure"] = "document"
                result["pageCount"] = self._count_docx_pages(content)
                raw_text = self._extract_docx_text(content)
            elif ext in {".txt", ".json", ".xml", ".md"}:
                result["structure"] = "text"
                result["pageCount"] = 1
                raw_text = content.decode("utf-8", errors="replace")
        except Exception as exc:
            logger.warning("document_parse_failed", filename=filename, error=str(exc))

        if raw_text:
            result["hasText"] = True
            result["textContent"] = raw_text[:50000]
            words = re.findall(r"\S+", raw_text)
            result["wordCount"] = len(words)
            result["languageHint"] = self._detect_language_hint(raw_text)

        return result

    def _count_pdf_pages(self, content: bytes) -> int:
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(content)
            return len(reader.pages)
        except Exception:
            return 0

    def _extract_pdf_text(self, content: bytes) -> str:
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(content)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            return ""

    def _count_docx_pages(self, content: bytes) -> int:
        try:
            from io import BytesIO
            import docx
            doc = docx.Document(BytesIO(content))
            return len(doc.paragraphs) // 40 + 1
        except Exception:
            return 0

    def _extract_docx_text(self, content: bytes) -> str:
        try:
            from io import BytesIO
            import docx
            doc = docx.Document(BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return ""

    def _extract_spreadsheet_text(self, content: bytes, ext: str) -> str:
        try:
            if ext == ".csv":
                return content.decode("utf-8", errors="replace")
            from io import BytesIO
            import openpyxl
            wb = openpyxl.load_workbook(BytesIO(content), read_only=True, data_only=True)
            lines: list[str] = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                for row in ws.iter_rows(values_only=True):
                    lines.append(",".join(str(cell) if cell is not None else "" for cell in row))
            return "\n".join(lines)
        except Exception:
            return ""

    def _detect_language_hint(self, text: str) -> str:
        sample = text[:2000].lower()
        spanish_patterns = ["factura", "proveedor", "total", "importe", "fecha", "cliente", "iva", "nif", "dirección", "teléfono"]
        score = sum(1 for pattern in spanish_patterns if pattern in sample)
        if score >= 2:
            return "es"
        return "en"

    def _extract_features(self, metadata: dict, parsed: dict, filename: str) -> dict[str, Any]:
        return {
            "filename": filename,
            "extension": metadata.get("extension", ""),
            "mimeType": metadata.get("mimeType", ""),
            "sizeBytes": metadata.get("sizeBytes", 0),
            "pageCount": parsed.get("pageCount", 0),
            "wordCount": parsed.get("wordCount", 0),
            "hasText": parsed.get("hasText", False),
            "isImage": metadata.get("isImage", False),
            "isPdf": metadata.get("isPdf", False),
            "isSpreadsheet": metadata.get("isSpreadsheet", False),
            "structure": parsed.get("structure", "unknown"),
            "languageHint": parsed.get("languageHint"),
            "extractedAt": self._now(),
        }
