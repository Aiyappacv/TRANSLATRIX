"""
DOCX Parser Adapter
Extract text and tables from Microsoft Word documents
"""
from pathlib import Path
from typing import List, Dict, Any
import structlog

from .base import BaseExtractor, ExtractionResult, ExtractionError

logger = structlog.get_logger(__name__)


class DOCXExtractor(BaseExtractor):
    """
    DOCX content extractor using python-docx
    Extracts paragraphs, tables, and document metadata
    """

    def __init__(self):
        self.supported_mimes = [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ]

    def can_extract(self, file_path: Path, mime_type: str) -> bool:
        """Check if this is a DOCX file"""
        return mime_type in self.supported_mimes or file_path.suffix.lower() in [".docx", ".doc"]

    def extract(
        self,
        file_path: Path,
        extract_tables: bool = True,
        extract_metadata: bool = True,
        **kwargs
    ) -> ExtractionResult:
        """
        Extract content from DOCX file

        Args:
            file_path: Path to DOCX file
            extract_tables: Whether to extract tables
            extract_metadata: Whether to extract metadata

        Returns:
            ExtractionResult with extracted content
        """
        try:
            from docx import Document

            logger.info("extracting_docx", file_path=str(file_path))

            doc = Document(file_path)

            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            full_text = "\n\n".join(text_parts)

            # Extract tables
            tables = []
            if extract_tables:
                for table_idx, table in enumerate(doc.tables):
                    rows_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        rows_data.append(row_data)

                    if rows_data:
                        headers = rows_data[0] if rows_data else []
                        data_rows = rows_data[1:] if len(rows_data) > 1 else []

                        tables.append({
                            "table_index": table_idx,
                            "headers": headers,
                            "rows": data_rows,
                        })

            # Extract metadata
            metadata = {}
            if extract_metadata:
                try:
                    core_props = doc.core_properties
                    metadata = {
                        "author": core_props.author or "",
                        "title": core_props.title or "",
                        "subject": core_props.subject or "",
                        "created": str(core_props.created) if core_props.created else "",
                        "modified": str(core_props.modified) if core_props.modified else "",
                        "last_modified_by": core_props.last_modified_by or "",
                    }
                except Exception as e:
                    logger.warning("metadata_extraction_failed", error=str(e))

            # DOCX native text extraction is reliable
            confidence = 0.98 if full_text.strip() else 0.0

            logger.info(
                "docx_extraction_complete",
                text_length=len(full_text),
                tables=len(tables),
                confidence=confidence,
            )

            return ExtractionResult(
                text=full_text,
                tables=tables,
                metadata=metadata,
                page_count=None,  # DOCX doesn't have explicit pages
                confidence=confidence,
            )

        except ImportError as e:
            raise ExtractionError(f"Required DOCX library not installed: {str(e)}")
        except Exception as e:
            logger.error("docx_extraction_error", error=str(e), file_path=str(file_path))
            raise ExtractionError(f"Failed to extract DOCX content: {str(e)}")

    def get_supported_formats(self) -> List[str]:
        """Return supported MIME types"""
        return self.supported_mimes
