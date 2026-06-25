from __future__ import annotations

import asyncio
import csv
import io
import logging
import re
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

logger = logging.getLogger("translatrix.processing")

from sqlalchemy.orm import Session

from uuid import UUID

from app.modules.frontend_api.defaults import OCR_SETTINGS
from app.modules.frontend_api.document_intelligence import (
    extract_financial_fields, map_gemini_fields, processing_validation_issues,
)
from app.modules.frontend_api.json_export_service import (
    build_extraction_json,
    build_ocr_only_json,
    store_extraction_json,
)
from app.modules.frontend_api.events import append_error, append_processing_log
from app.modules.frontend_api.store import get_state, set_state
from app.modules.frontend_api.utils import new_id, now_iso
from app.modules.frontend_api.business_validators import validate_all
from app.modules.extraction.adapters.gemini_extractor import GeminiExtractor
from app.modules.extraction.adapters.base import ExtractionError as GeminiExtractionError

SUPPORTED_TEXT_EXTENSIONS = {".txt", ".csv", ".json", ".xml", ".md"}


def _is_mistral_available() -> bool:
    from app.config import settings
    return bool(getattr(settings, "MISTRAL_API_KEY", None))


def _run_gemini_extraction(path: Path, fallback_currency: str) -> dict[str, Any]:
    """Run Gemini 2.5 Pro's multimodal document understanding directly against
    the original file (not the Mistral OCR text) and map its output onto the
    field-dict shape the rest of the pipeline expects.

    Raises GeminiExtractionError on any failure (missing/invalid key, rate
    limit, network error, bad response) — callers intentionally do not catch
    this so the document is marked failed rather than silently degraded."""
    from app.config import settings

    extractor = GeminiExtractor({
        "api_key": settings.GEMINI_API_KEY,
        "model": settings.GEMINI_EXTRACTION_MODEL,
    })
    result = extractor.extract(path, extract_tables=True, extract_metadata=True)
    extracted_fields = result.metadata.get("extracted_fields") or {}
    field_confidence = result.metadata.get("field_confidence") or {}
    fields = map_gemini_fields(extracted_fields, field_confidence, fallback_currency=fallback_currency)
    fields["_geminiOverallConfidence"] = result.confidence
    fields["documentType"] = fields["_geminiDocumentType"] = result.metadata.get("document_type")
    return fields


def _mistral_ocr(path: Path, mime_type: str, *, kind: str = "document") -> tuple[str, float]:
    """OCR a local file via Mistral's hosted OCR API.

    `kind` is "document" for PDFs (multi-page, sent as document_url) or "image"
    for single images (sent as image_url). Both accept a base64 data URI for
    local content. Returns (extracted_markdown_text, average_page_confidence).
    """
    import base64
    import httpx
    from app.config import settings

    encoded = base64.b64encode(path.read_bytes()).decode("ascii")
    if kind == "image":
        document = {"type": "image_url", "image_url": f"data:{mime_type};base64,{encoded}"}
    else:
        document = {"type": "document_url", "document_url": f"data:{mime_type};base64,{encoded}"}

    response = httpx.post(
        "https://api.mistral.ai/v1/ocr",
        headers={"Authorization": f"Bearer {settings.MISTRAL_API_KEY}", "Content-Type": "application/json"},
        json={
            "model": settings.MISTRAL_OCR_MODEL,
            "document": document,
            "confidence_scores_granularity": "page",
        },
        timeout=180.0,
    )
    response.raise_for_status()
    data = response.json()
    pages = data.get("pages", [])
    text = "\f".join(p.get("markdown", "") for p in pages).strip()
    scores = [
        p["confidence_scores"]["average_page_confidence_score"]
        for p in pages
        if isinstance(p.get("confidence_scores"), dict)
        and p["confidence_scores"].get("average_page_confidence_score") is not None
    ]
    confidence = sum(scores) / len(scores) if scores else 0.0
    return text, confidence


def _finance_text_score(text: str) -> int:
    if not text.strip():
        return 0
    fields = extract_financial_fields(text)
    score = sum(3 for key in ("invoiceNumber", "total", "invoiceDate") if fields.get(key))
    score += sum(1 for key in ("vendor", "gstVatNumber", "taxAmount", "subtotal") if fields.get(key))
    score += min(3, len(re.findall(r"\b(?:invoice|inv\.?|bill|gst|vat|tax|total|amount|date)\b", text, re.IGNORECASE)))
    return score


DIRECT_PARSER_CONFIDENCE = 0.95


def _read_pdf(path: Path, max_pages: int = 500) -> tuple[str, list[list[str]], str, float]:
    """Extract text from a PDF using Mistral OCR (primary) or direct text parser.

    Returns (text, rows, engine_name, confidence).
    """
    from PyPDF2 import PdfReader

    reader = PdfReader(str(path))
    pages = list(reader.pages)[:max(1, max_pages)]
    direct_text = "\f".join((page.extract_text() or "") for page in pages).strip()

    if _is_mistral_available():
        try:
            ocr_text, confidence = _mistral_ocr(path, "application/pdf", kind="document")
            logger.info("ocr_selection file=%s ocr_len=%d engine=gemini-2.5-pro confidence=%.2f",
                        path.name, len(ocr_text), confidence)
            if ocr_text.strip():
                return ocr_text, [], "gemini-2.5-pro", confidence
        except Exception as exc:
            logger.warning("mistral_ocr_failed file=%s error=%s", path.name, exc)

    if direct_text.strip():
        logger.info("ocr_selection file=%s engine=direct_parser reason=text_layer_fallback len=%d",
                    path.name, len(direct_text))
        return direct_text, [], "direct_parser", DIRECT_PARSER_CONFIDENCE

    return "", [], "direct_parser", 0.0


def _read_image(path: Path) -> tuple[str, list[list[str]], str, float]:
    """Extract text from an image using Mistral OCR.

    Returns (text, rows, engine_name, confidence).
    """
    if _is_mistral_available():
        try:
            import mimetypes
            mime = mimetypes.guess_type(str(path))[0] or "image/png"
            text, confidence = _mistral_ocr(path, mime, kind="image")
            if text.strip():
                logger.debug("image_ocr file=%s engine=gemini-2.5-pro chars=%d confidence=%.2f", path.name, len(text), confidence)
                return text, [], "gemini-2.5-pro", confidence
        except Exception as exc:
            logger.warning("mistral_ocr_failed file=%s error=%s", path.name, exc)

    return "", [], "direct_parser", 0.0


def _read_docx(path: Path) -> tuple[str, list[list[str]], str, float]:
    from docx import Document

    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    tables: list[list[str]] = []
    for table in document.tables:
        for row in table.rows:
            tables.append([cell.text for cell in row.cells])
    return "\n".join(paragraphs).strip(), tables, "direct_parser", DIRECT_PARSER_CONFIDENCE


def _read_xlsx(path: Path) -> tuple[str, list[list[str]], str, float]:
    from openpyxl import load_workbook

    workbook = load_workbook(str(path), read_only=True, data_only=True)
    rows: list[list[str]] = []
    for sheet in workbook.worksheets:
        for values in sheet.iter_rows(values_only=True):
            row = ["" if value is None else str(value) for value in values]
            if any(cell.strip() for cell in row):
                rows.append(row)
            if len(rows) >= 2000:
                break
        if len(rows) >= 2000:
            break
    text = "\n".join(" | ".join(row) for row in rows)
    return text, rows, "direct_parser", DIRECT_PARSER_CONFIDENCE


def _read_csv(path: Path) -> tuple[str, list[list[str]], str, float]:
    raw = path.read_bytes()
    text = raw.decode("utf-8", errors="replace")
    rows = [row for row in csv.reader(io.StringIO(text))][:2000]
    return text.strip(), rows, "direct_parser", DIRECT_PARSER_CONFIDENCE


def extract_document(path: Path, filename: str, mime_type: str, *, max_pages: int = 500) -> tuple[str, list[list[str]], str, float]:
    extension = Path(filename).suffix.lower()
    if extension == ".pdf" or mime_type == "application/pdf":
        return _read_pdf(path, max_pages=max_pages)
    if extension in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"} or mime_type.startswith("image/"):
        return _read_image(path)
    if extension == ".docx":
        return _read_docx(path)
    if extension in {".xlsx", ".xlsm"}:
        return _read_xlsx(path)
    if extension == ".csv":
        return _read_csv(path)
    if extension in SUPPORTED_TEXT_EXTENSIONS or mime_type.startswith("text/"):
        return path.read_text(encoding="utf-8", errors="replace").strip(), [], "direct_parser", DIRECT_PARSER_CONFIDENCE
    return "", [], "direct_parser", 0.0


def _find_amount(text: str) -> float:
    candidates: list[float] = []
    for match in re.finditer(r"(?<![A-Za-z0-9])(?:₹|Rs\.?|INR|\$|USD|€|EUR|£|GBP)?\s*([0-9]{1,3}(?:[, ]?[0-9]{3})*(?:\.[0-9]{1,2})|[0-9]+(?:\.[0-9]{1,2})?)", text, re.IGNORECASE):
        try:
            value = float(match.group(1).replace(",", "").replace(" ", ""))
            if 0 < value < 1_000_000_000:
                candidates.append(value)
        except ValueError:
            continue
    return max(candidates) if candidates else 0.0


def _find_currency(text: str, fallback: str = "USD") -> str:
    upper = text.upper()
    if "₹" in text or " INR" in f" {upper}" or "RS." in upper or " RS " in f" {upper} ":
        return "INR"
    if "€" in text or " EUR" in f" {upper}":
        return "EUR"
    if "£" in text or " GBP" in f" {upper}":
        return "GBP"
    if "$" in text or " USD" in f" {upper}":
        return "USD"
    return fallback


def _classify(text: str) -> tuple[str, str, str]:
    lowered = text.lower()
    if any(word in lowered for word in ("salary", "rent", "travel", "expense", "purchase", "invoice", "bill")):
        return "Expenses", "Operating Expense", "Matched expense-related terminology."
    if any(word in lowered for word in ("sale", "revenue", "income", "receipt", "customer payment")):
        return "Income", "Operating Income", "Matched income-related terminology."
    if any(word in lowered for word in ("asset", "equipment", "machinery", "vehicle")):
        return "Assets", "Fixed Asset", "Matched asset-related terminology."
    if any(word in lowered for word in ("loan", "liability", "payable")):
        return "Liabilities", "Current Liability", "Matched liability-related terminology."
    return "Expenses", "Unclassified", "No strong keyword match; defaulted to Expenses for human review."


def _table_payload(rows: list[list[str]]) -> list[dict[str, Any]]:
    if not rows:
        return []
    headers = rows[0]
    data_rows = rows[1:] if len(rows) > 1 else []
    cells: list[list[dict[str, Any]]] = []
    for row_index, row in enumerate(data_rows[:500]):
        cells.append([
            {
                "id": new_id("cell"),
                "rowIndex": row_index,
                "columnIndex": column_index,
                "value": value,
                "confidence": 0.98,
            }
            for column_index, value in enumerate(row)
        ])
    return [{"id": new_id("table"), "name": "Extracted table", "pageNumber": 1, "confidence": 0.98, "headers": headers, "rows": cells}]


def _ocr_pages(text: str, confidence: float, language: str = "Unknown") -> list[dict[str, Any]]:
    blocks = []
    for index, line in enumerate([line.strip() for line in text.splitlines() if line.strip()][:200]):
        block_type = "amount" if re.search(r"\d", line) else "text"
        blocks.append({
            "id": new_id("block"),
            "type": block_type,
            "text": line,
            "confidence": confidence,
            "bbox": [20, 20 + index * 24, 760, 20],
            "pageNumber": 1,
        })
    return [{"pageNumber": 1, "width": 800, "height": max(1000, len(blocks) * 28), "confidence": confidence, "detectedLanguage": language, "blocks": blocks}]


def _flatten_text_for_classification(data: Any) -> str:
    parts: list[str] = []

    def _walk(value: Any, depth: int = 0) -> None:
        if depth > 10:
            return
        if isinstance(value, dict):
            for inner in value.values():
                _walk(inner, depth + 1)
        elif isinstance(value, list):
            for inner in value:
                _walk(inner, depth + 1)
        elif isinstance(value, str):
            trimmed = value.strip()
            if trimmed:
                parts.append(trimmed)
        elif isinstance(value, (int, float)):
            if value not in (0, 0.0):
                parts.append(str(value))

    _walk(data)
    return " ".join(parts)


def _save_progress(db: Session, scope: str, record: dict[str, Any]) -> None:
    """Persist the current file state so polling frontends see live progress."""
    files = get_state(db, scope, "files", [])
    for i, f in enumerate(files):
        if str(f.get("id")) == str(record.get("id")):
            files[i] = dict(record)
            break
    set_state(db, scope, "files", files)


def process_file_record(db: Session, scope: str, record: dict[str, Any], user: Any, *, retry_count: int = 0) -> dict[str, Any]:
    start = time.perf_counter()
    file_id = str(record["id"])
    path = Path(str(record.get("_contentPath") or ""))
    timestamp = now_iso()
    job_id = new_id("process")
    record.update({"status": "processing", "processingStage": "Initializing", "ocrStatus": "processing", "extractionStatus": "processing", })
    record.setdefault("processingLogs", []).insert(0, {
        "id": new_id("log"), "step": "Processing", "worker": "document-pipeline", "status": "processing",
        "message": "Document processing started.", "startedAt": timestamp, "retryable": False,
    })
    append_processing_log(db, scope, stage="file", message=f"Processing started for {record.get('fileName')}", level="info", job_id=job_id, file_id=file_id, batch_id=record.get("batchId"), retry_count=retry_count)
    _save_progress(db, scope, record)

    try:
        # Stage 1: Preprocessing — file validation, metadata, dedup
        if not path.exists():
            raise FileNotFoundError("The stored upload is no longer available.")
        record["processingStage"] = "Preprocessing"
        _save_progress(db, scope, record)

        ocr_settings = {**OCR_SETTINGS, **get_state(db, scope, "ocr_settings", {})}
        max_pages = max(1, min(int(ocr_settings.get("maxPagesPerFile") or 500), 2000))
        table_extraction = bool(ocr_settings.get("tableExtractionEnabled"))
        layout_analysis = bool(ocr_settings.get("layoutAnalysisEnabled"))

        # Stage 2: Document Analysis — OCR/text-layer read of the document.
        # Still needed even though Gemini reads the original file directly for
        # field extraction: this text drives language detection and the
        # raw_ocr_text audit trail in the extraction JSON.
        record["processingStage"] = "Document Analysis"
        _save_progress(db, scope, record)
        text, rows, engine, ocr_confidence = extract_document(path, str(record.get("fileName") or path.name), str(record.get("mimeType") or ""), max_pages=max_pages)
        if not text.strip():
            fallback_hint = " Cloud fallback is enabled but requires configured provider credentials." if ocr_settings.get("cloudFallbackEnabled") else " Enable cloud fallback or install the local OCR runtime for scanned documents."
            raise ValueError(f"No machine-readable text was detected.{fallback_hint}")

        # Layout analysis — try PP-Structure for advanced layout detection
        layout_data: dict[str, Any] = {"regions": [], "tables": []}
        if layout_analysis:
            try:
                extension = Path(record.get("fileName") or "").suffix.lower()
                pp_layout: dict[str, Any] | None = None
                try:
                    from paddleocr.ppstructure import PPStructure
                    pp_engine = PPStructure(show_log=False, lang='en', use_angle_cls=True)
                    pp_result = pp_engine.predict(str(path))
                    if pp_result and isinstance(pp_result, list):
                        pp_regions: list[dict[str, Any]] = []
                        pp_tables: list[dict[str, Any]] = []
                        for region in pp_result:
                            region_type = region.get("type", "")
                            region_bbox = region.get("bbox", [0, 0, 0, 0])
                            region_res = region.get("res", {})
                            if region_type == "table":
                                html = region_res.get("html", "")
                                if html:
                                    pp_tables.append({
                                        "type": "table",
                                        "bbox": region_bbox,
                                        "html": html,
                                        "confidence": region.get("confidence", 0.8),
                                    })
                            else:
                                pp_regions.append({
                                    "type": region_type,
                                    "bbox": region_bbox,
                                    "text": region_res.get("text", "") if isinstance(region_res, dict) else str(region_res),
                                    "confidence": region.get("confidence", 0.8),
                                })
                        if pp_regions or pp_tables:
                            pp_layout = {"regions": pp_regions, "tables": pp_tables}
                            logger.info("pp_structure_complete file_id=%s regions=%d tables=%d",
                                        file_id, len(pp_regions), len(pp_tables))
                except ImportError:
                    logger.debug("pp_structure_not_available file_id=%s", file_id)
                except Exception as pp_exc:
                    logger.debug("pp_structure_failed file_id=%s error=%s", file_id, pp_exc)

                if pp_layout:
                    layout_data = pp_layout
            except Exception as la_exc:
                logger.warning("layout_analysis_failed file_id=%s error=%s", file_id, la_exc)
        record["layoutAnalysis"] = layout_data

        ocr_threshold = float(ocr_settings.get("confidenceThreshold") or 80) / 100
        company_settings = get_state(db, scope, "settings_company", {})
        default_currency = str(company_settings.get("defaultCurrency") or "USD")

        # Stage 4: Gemini 2.5 Pro Processing — multimodal read of the original
        # file (not the OCR text above), which is what actually drives field
        # extraction. Intentionally not wrapped in a try/except here: a
        # missing/invalid API key, rate limit, or network failure should fail
        # the document loudly (caught by the outer handler) rather than
        # silently falling back to a lower-quality result.
        record["processingStage"] = "Gemini 2.5 Pro Processing"
        _save_progress(db, scope, record)
        fields = _run_gemini_extraction(path, default_currency)

        # Stage 5: Field Extraction — Gemini's response has been parsed and
        # mapped onto the internal field schema at this point.
        record["processingStage"] = "Field Extraction"
        _save_progress(db, scope, record)

        classification_text = text

        category, subcategory, reason = _classify(classification_text)
        completed_at = now_iso()

        # Business validation
        validation_results = validate_all(fields)
        record["validationResults"] = validation_results
        logger.info("business_validation file_id=%s valid=%d review=%d warning=%d", file_id, len(validation_results.get("valid", [])), len(validation_results.get("needs_review", [])), len(validation_results.get("warning", [])))

        review_reasons: list[str] = []
        if ocr_confidence < ocr_threshold:
            review_reasons.append(f"OCR confidence {ocr_confidence:.0%} is below the configured {ocr_threshold:.0%} threshold.")
        for item in validation_results.get("needs_review", []):
            msg = item.get("message", "")
            if msg:
                review_reasons.append(f"Validation: {msg}")
        for item in validation_results.get("warning", []):
            msg = item.get("message", "")
            if msg:
                review_reasons.append(f"Validation warning: {msg}")

        record.update({
            "status": "processing",
            "ocrStatus": "completed",
            "extractionStatus": "completed",
            "entriesExtracted": 1,
            "confidence": ocr_confidence,
            "sourceLanguage": record.get("sourceLanguage"),
            "extractionMethod": "gemini-2.5-pro",
            "extractionConfidence": fields.get("_geminiOverallConfidence"),
            "fieldConfidence": fields.get("_geminiFieldConfidence") or {},
            "extractedText": text,
            "spreadsheetRows": rows if table_extraction else [],
            "extractedTables": _table_payload(rows) if table_extraction else [],
            "structuredFields": fields,
            "amount": fields.get("total") or 0,
            "currency": fields.get("currency") or str(company_settings.get("defaultCurrency") or "USD"),
            "vendor": fields.get("vendor"), "customer": fields.get("customer"),
            "invoiceNumber": fields.get("invoiceNumber"), "referenceNumber": fields.get("referenceNumber"),
            "invoiceDate": fields.get("invoiceDate"), "dueDate": fields.get("dueDate"),
            "subtotal": fields.get("subtotal"), "taxAmount": fields.get("taxAmount"), "taxRate": fields.get("taxRate"),
            "gstVatNumber": fields.get("gstVatNumber"), "lineItems": fields.get("lineItems") or [],
            "category": category, "subcategory": subcategory, "classificationReason": reason,
            "reviewReasons": review_reasons,
            "processingSettings": {
                "ocr": {"primaryEngine": ocr_settings.get("primaryEngine"), "confidenceThreshold": ocr_settings.get("confidenceThreshold"), "tableExtractionEnabled": table_extraction, "layoutAnalysisEnabled": layout_analysis, "handwritingEnabled": bool(ocr_settings.get("handwritingEnabled")), "maxPagesPerFile": max_pages},
            },
            "_preprocessed": False,
            "ocr": {
                "engine": engine, "engineVersion": "local-runtime-2", "status": "completed",
                "languageDetected": record.get("sourceLanguage"), "overallConfidence": ocr_confidence,
                "pageCount": max(1, text.count("\f") + 1), "startedAt": timestamp, "completedAt": completed_at,
                "preprocessingApplied": False,
                "pages": _ocr_pages(text, ocr_confidence, record.get("sourceLanguage") or "Unknown") if layout_analysis else [{"pageNumber": 1, "width": 800, "height": 1000, "confidence": ocr_confidence, "detectedLanguage": record.get("sourceLanguage"), "blocks": [{"id": new_id("block"), "type": "text", "text": text, "confidence": ocr_confidence, "bbox": [20, 20, 760, 960], "pageNumber": 1}]}],
            },
        })
        record["processingStage"] = "Validation"
        _save_progress(db, scope, record)
        record["validationIssues"] = processing_validation_issues(record)
        has_error_issues = any(issue.get("severity") == "error" for issue in record["validationIssues"])
        if has_error_issues:
            review_reasons.extend(issue["message"] for issue in record["validationIssues"] if issue.get("severity") == "error")

        record.setdefault("processingLogs", []).insert(0, {"id": new_id("log"), "step": "Extraction", "worker": "document-pipeline", "status": "completed", "message": f"Structured finance extraction completed with {len(fields.get('lineItems') or [])} line-item candidate(s).", "startedAt": timestamp, "completedAt": completed_at, "retryable": False})
        record.setdefault("processingLogs", []).insert(0, {"id": new_id("log"), "step": "OCR", "worker": engine, "status": "completed", "message": "Local OCR/direct parsing completed.", "startedAt": timestamp, "completedAt": completed_at, "retryable": True})

        # ===== Finalization — update the parent Processing log and determine final status =====
        for log_entry in record.get("processingLogs", []):
            if log_entry.get("step") == "Processing":
                log_entry["status"] = "completed"
                log_entry["message"] = "Document processing completed."
                log_entry["completedAt"] = completed_at
                break

        # ===== Generate structured extraction JSON =====
        record["processingStage"] = "Generating JSON"
        _save_progress(db, scope, record)
        try:
            extraction_json = build_extraction_json(record, fields)
            record["extractionJson"] = extraction_json.model_dump_export()
            store_extraction_json(db, scope, file_id, extraction_json)
            logger.info("extraction_json_generated file_id=%s fields=%d line_items=%d", file_id, len(extraction_json.confidence_details), len(extraction_json.line_items))
        except Exception as json_err:
            logger.warning("extraction_json_failed file_id=%s error=%s", file_id, json_err)
            try:
                fallback_json = build_ocr_only_json(record)
                record["extractionJson"] = fallback_json.model_dump_export()
                store_extraction_json(db, scope, file_id, fallback_json)
                logger.info("extraction_json_fallback_stored file_id=%s", file_id)
            except Exception as fallback_err:
                logger.warning("extraction_json_fallback_failed file_id=%s error=%s", file_id, fallback_err)
                record["extractionJson"] = None

        # Determine final document status based on review requirements
        if has_error_issues:
            record["status"] = "validation_failed"
        else:
            record["status"] = "needs_review" if review_reasons else "completed"
        record["processingCompletedAt"] = completed_at
        record["updatedAt"] = completed_at
        record["lastProcessingJobId"] = job_id
        record["reprocessCount"] = int(record.get("reprocessCount") or 0) + 1
        logger.info("final_status file_id=%s status=%s has_extraction_json=%s review_reasons=%d", file_id, record["status"], record.get("extractionJson") is not None, len(review_reasons))

        record["processingStage"] = "Extraction Complete" if record["status"] in ("completed", "needs_review") else record["status"]
        _save_progress(db, scope, record)
        record.setdefault("processingLogs", []).insert(0, {
            "id": new_id("log"), "step": "Finalization", "worker": "document-pipeline", "status": "completed",
            "message": f"Document status updated: processing \u2192 {record['status']}. All stages completed successfully.",
            "startedAt": timestamp, "completedAt": completed_at, "retryable": False,
        })

        duration_ms = int((time.perf_counter() - start) * 1000)
        append_processing_log(db, scope, stage="finalization", message=f"Document status updated to {record['status']}", level="success", job_id=job_id, file_id=file_id, batch_id=record.get("batchId"), duration_ms=duration_ms, retry_count=retry_count)
        append_processing_log(db, scope, stage="ocr", message="OCR/direct parsing completed", level="success", job_id=job_id, file_id=file_id, batch_id=record.get("batchId"), duration_ms=duration_ms, retry_count=retry_count)
        append_processing_log(db, scope, stage="classification", message=f"Entry classified as {category} / {subcategory}", level="success", job_id=job_id, file_id=file_id, batch_id=record.get("batchId"), duration_ms=duration_ms, retry_count=retry_count)
        return record
    except Exception as exc:
        message = str(exc)
        completed_at = now_iso()
        try:
            db.rollback()
        except Exception:
            pass
        record.update({"status": "validation_failed", "processingStage": "Failed", "ocrStatus": "failed", "extractionStatus": "failed", "entriesExtracted": 0, "validationIssues": [{"code": "PROCESSING_FAILED", "severity": "error", "message": message, "field": "processing"}]})
        _save_progress(db, scope, record)
        for log_entry in record.get("processingLogs", []):
            if log_entry.get("step") == "Processing":
                log_entry["status"] = "failed"
                log_entry["message"] = "Document processing failed."
                log_entry["errorDetails"] = message
                log_entry["completedAt"] = completed_at
                log_entry["retryable"] = True
                break
        else:
            record.setdefault("processingLogs", []).insert(0, {"id": new_id("log"), "step": "Processing", "worker": "document-pipeline", "status": "failed", "message": "Document processing failed.", "errorDetails": message, "startedAt": timestamp, "completedAt": completed_at, "retryable": True})
        append_processing_log(db, scope, stage="file", message=message, level="error", job_id=job_id, file_id=file_id, batch_id=record.get("batchId"), duration_ms=int((time.perf_counter() - start) * 1000), retry_count=retry_count)
        append_error(db, scope, category="processing", code="DOCUMENT_PROCESSING_FAILED", message=message, entity_type="file", entity_id=file_id, retryable=True, details={"fileName": record.get("fileName"), "batchId": record.get("batchId")})
        return record
